"""
Report generator module for creating PDF/DOCX reports
"""

from pathlib import Path
from typing import List, Dict, Optional
from datetime import datetime
import config
from .utils import ensure_dir, format_timestamp_readable, get_file_size_mb


def generate_report(
    video_path: str,
    transcription_result: Dict,
    scene_changes: List[Dict],
    clicks: List[Dict],
    output_path: Optional[str] = None
) -> Optional[str]:
    """
    Generate meeting report in DOCX or PDF format
    
    Args:
        video_path: Path to original video file
        transcription_result: Whisper transcription result
        scene_changes: List of detected scene changes
        clicks: List of detected UI interactions
        output_path: Optional output path for report
        
    Returns:
        Path to generated report file or None if error
    """
    try:
        video_path = Path(video_path)
        
        if output_path is None:
            ensure_dir(config.REPORTS_DIR)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            extension = config.REPORT_FORMAT
            output_path = config.REPORTS_DIR / f"meeting_report_{timestamp}.{extension}"
        else:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
        
        print(f"Generating {config.REPORT_FORMAT.upper()} report: {output_path}")
        
        if config.REPORT_FORMAT.lower() == "docx":
            return generate_docx_report(
                video_path, transcription_result, scene_changes, clicks, output_path
            )
        elif config.REPORT_FORMAT.lower() == "pdf":
            return generate_pdf_report(
                video_path, transcription_result, scene_changes, clicks, output_path
            )
        else:
            print(f"Error: Unsupported report format: {config.REPORT_FORMAT}")
            return None
            
    except Exception as e:
        print(f"Error generating report: {e}")
        return None


def generate_docx_report(
    video_path: Path,
    transcription_result: Dict,
    scene_changes: List[Dict],
    clicks: List[Dict],
    output_path: Path
) -> Optional[str]:
    """
    Generate DOCX report using python-docx
    
    Args:
        video_path: Path to video file
        transcription_result: Transcription result
        scene_changes: Scene changes list
        clicks: Clicks list
        output_path: Output file path
        
    Returns:
        Path to generated report or None if error
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Meeting Video Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Video Information
        doc.add_heading('Video Information', level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('Video File', video_path.name),
            ('File Size', f"{get_file_size_mb(str(video_path)):.2f} MB"),
            ('Report Generated', datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ('Total Duration', format_timestamp_readable(transcription_result.get('segments', [{}])[-1].get('end', 0) if transcription_result.get('segments') else 0))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = str(value)
        
        # Scene Changes
        if config.INCLUDE_SCREENSHOTS and scene_changes:
            doc.add_heading('Scene Changes & Screenshots', level=1)
            doc.add_paragraph(f'Total scene changes detected: {len(scene_changes)}')
            
            for scene in scene_changes:
                doc.add_heading(f"Scene {scene['index'] + 1} - {scene['timestamp_readable']}", level=2)
                
                # Add screenshot if available
                frame_path = Path(scene['frame_path'])
                if frame_path.exists():
                    try:
                        doc.add_picture(str(frame_path), width=Inches(5))
                    except Exception as e:
                        doc.add_paragraph(f"[Screenshot unavailable: {e}]")
                
                doc.add_paragraph(f"Timestamp: {scene['timestamp_readable']}")
                doc.add_paragraph(f"Similarity Score: {scene['similarity']:.3f}")
                doc.add_paragraph("")
        
        # Transcript
        if config.INCLUDE_TRANSCRIPT:
            doc.add_heading('Full Transcript', level=1)
            full_text = transcription_result.get('text', '').strip()
            if full_text:
                doc.add_paragraph(full_text)
            else:
                doc.add_paragraph("No transcript available.")
        
        # Segmented Transcript
        doc.add_heading('Segmented Transcript with Timestamps', level=1)
        segments = transcription_result.get('segments', [])
        
        if segments:
            for segment in segments:
                timestamp = format_timestamp_readable(segment['start'])
                text = segment['text'].strip()
                doc.add_paragraph(f"[{timestamp}] {text}")
        else:
            doc.add_paragraph("No segments available.")
        
        # UI Interactions
        if config.INCLUDE_CLICK_LOG and clicks:
            doc.add_heading('UI Interactions (Clicks & Movements)', level=1)
            doc.add_paragraph(f'Total interactions detected: {len(clicks)}')
            
            click_table = doc.add_table(rows=len(clicks) + 1, cols=2)
            click_table.style = 'Light Grid Accent 1'
            
            # Header
            click_table.rows[0].cells[0].text = 'Timestamp'
            click_table.rows[0].cells[1].text = 'Intensity'
            
            # Data
            for i, click in enumerate(clicks, 1):
                click_table.rows[i].cells[0].text = click['timestamp_readable']
                click_table.rows[i].cells[1].text = f"{click['intensity']:.2f}"
        
        # Summary
        doc.add_heading('Summary', level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"This meeting video contains {len(scene_changes)} major scene changes and "
                           f"{len(clicks)} UI interactions. The total duration is approximately "
                           f"{format_timestamp_readable(segments[-1].get('end', 0) if segments else 0)}.")
        
        # Save document
        doc.save(str(output_path))
        print(f"DOCX report generated: {output_path}")
        return str(output_path)
        
    except ImportError:
        print("Error: python-docx not installed. Install it with: pip install python-docx")
        return None
    except Exception as e:
        print(f"Error generating DOCX report: {e}")
        return None


def generate_pdf_report(
    video_path: Path,
    transcription_result: Dict,
    scene_changes: List[Dict],
    clicks: List[Dict],
    output_path: Path
) -> Optional[str]:
    """
    Generate PDF report using ReportLab
    
    Args:
        video_path: Path to video file
        transcription_result: Transcription result
        scene_changes: Scene changes list
        clicks: Clicks list
        output_path: Output file path
        
    Returns:
        Path to generated report or None if error
    """
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        doc = SimpleDocTemplate(str(output_path), pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        story.append(Paragraph('Meeting Video Report', title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Video Information
        story.append(Paragraph('Video Information', styles['Heading1']))
        info_data = [
            ['Video File', video_path.name],
            ['File Size', f"{get_file_size_mb(str(video_path)):.2f} MB"],
            ['Report Generated', datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
            ['Total Duration', format_timestamp_readable(transcription_result.get('segments', [{}])[-1].get('end', 0) if transcription_result.get('segments') else 0)]
        ]
        
        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(info_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Scene Changes
        if config.INCLUDE_SCREENSHOTS and scene_changes:
            story.append(Paragraph('Scene Changes & Screenshots', styles['Heading1']))
            story.append(Paragraph(f'Total scene changes detected: {len(scene_changes)}', styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
            
            for scene in scene_changes:
                story.append(Paragraph(f"Scene {scene['index'] + 1} - {scene['timestamp_readable']}", styles['Heading2']))
                
                frame_path = Path(scene['frame_path'])
                if frame_path.exists():
                    try:
                        img = Image(str(frame_path), width=5*inch, height=3*inch)
                        story.append(img)
                    except Exception as e:
                        story.append(Paragraph(f"[Screenshot unavailable: {e}]", styles['Normal']))
                
                story.append(Paragraph(f"Timestamp: {scene['timestamp_readable']}", styles['Normal']))
                story.append(Paragraph(f"Similarity Score: {scene['similarity']:.3f}", styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
        
        # Transcript
        if config.INCLUDE_TRANSCRIPT:
            story.append(PageBreak())
            story.append(Paragraph('Full Transcript', styles['Heading1']))
            full_text = transcription_result.get('text', '').strip()
            if full_text:
                story.append(Paragraph(full_text, styles['Normal']))
            else:
                story.append(Paragraph("No transcript available.", styles['Normal']))
        
        # Segmented Transcript
        story.append(PageBreak())
        story.append(Paragraph('Segmented Transcript with Timestamps', styles['Heading1']))
        segments = transcription_result.get('segments', [])
        
        if segments:
            for segment in segments:
                timestamp = format_timestamp_readable(segment['start'])
                text = segment['text'].strip()
                story.append(Paragraph(f"[{timestamp}] {text}", styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        else:
            story.append(Paragraph("No segments available.", styles['Normal']))
        
        # UI Interactions
        if config.INCLUDE_CLICK_LOG and clicks:
            story.append(PageBreak())
            story.append(Paragraph('UI Interactions (Clicks & Movements)', styles['Heading1']))
            story.append(Paragraph(f'Total interactions detected: {len(clicks)}', styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
            
            click_data = [['Timestamp', 'Intensity']]
            for click in clicks:
                click_data.append([click['timestamp_readable'], f"{click['intensity']:.2f}"])
            
            click_table = Table(click_data, colWidths=[3*inch, 3*inch])
            click_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(click_table)
        
        # Summary
        story.append(PageBreak())
        story.append(Paragraph('Summary', styles['Heading1']))
        summary_text = (f"This meeting video contains {len(scene_changes)} major scene changes and "
                       f"{len(clicks)} UI interactions. The total duration is approximately "
                       f"{format_timestamp_readable(segments[-1].get('end', 0) if segments else 0)}.")
        story.append(Paragraph(summary_text, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        print(f"PDF report generated: {output_path}")
        return str(output_path)
        
    except ImportError:
        print("Error: ReportLab not installed. Install it with: pip install reportlab")
        return None
    except Exception as e:
        print(f"Error generating PDF report: {e}")
        return None

